#!/usr/bin/env python3
"""
Generate professional DOCX files from MedBasket security assessment deliverables.
Uses python-docx to create formatted Word documents suitable for board presentation.

Output:
  MedBasket_Security_Assessment_Report.docx
  MedBasket_PoC_Walkthrough.docx
  MedBasket_Findings_Evidence.docx
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(BASE_DIR)
OUT_DIR = os.path.join(BASE_DIR, "docx_output")

# Colors
RED = RGBColor(220, 38, 38)
DARK_RED = RGBColor(153, 27, 27)
ORANGE = RGBColor(234, 88, 12)
GREEN = RGBColor(22, 163, 74)
BLUE = RGBColor(37, 99, 235)
DARK_BLUE = RGBColor(30, 58, 138)
GRAY = RGBColor(100, 116, 139)
DARK = RGBColor(15, 23, 42)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_doc(doc):
    """Apply base styling to document."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Calibri"
        hstyle.font.color.rgb = DARK_BLUE if level == 1 else DARK
        hstyle.font.bold = True
        if level == 1:
            hstyle.font.size = Pt(22)
            hstyle.paragraph_format.space_before = Pt(24)
        elif level == 2:
            hstyle.font.size = Pt(16)
            hstyle.paragraph_format.space_before = Pt(18)
        else:
            hstyle.font.size = Pt(13)
            hstyle.paragraph_format.space_before = Pt(12)

    return doc


def add_cover_page(doc, title, subtitle, classification="CONFIDENTIAL"):
    """Add a professional cover page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph("")

    # Classification banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"  {classification}  ")
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.highlight_color = 6  # Red highlight

    doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    doc.add_paragraph("")

    # Metadata table
    meta = [
        ("Target", "api.medbasket.com / admin.medbasket.com"),
        ("Date", "2026-02-18"),
        ("Classification", classification),
        ("Composite Risk", "10.0 / 10.0 — CRITICAL"),
        ("Assessor", "[Your Name / Organization]"),
        ("Authorization", "[Insert authorization document ID]"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(meta):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = GRAY
        if "CRITICAL" in val:
            table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RED
            table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()


def add_severity_badge(paragraph, severity):
    """Add a colored severity indicator."""
    colors = {
        "CRITICAL": RED,
        "HIGH": ORANGE,
        "MEDIUM": RGBColor(202, 138, 4),
        "LOW": GREEN,
        "INFO": GRAY,
        "INFORMATIONAL": GRAY,
    }
    color = colors.get(severity.upper(), GRAY)
    run = paragraph.add_run(f" [{severity.upper()}] ")
    run.font.color.rgb = color
    run.font.bold = True
    run.font.size = Pt(10)


def add_code_block(doc, code, language=""):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)


def add_finding_table(doc, finding_id, title, severity, cvss, owasp, endpoint, auth, records=""):
    """Add a standardized finding summary table."""
    data = [
        ("Finding", f"{finding_id}: {title}"),
        ("Severity", severity),
        ("CVSS v3.1", cvss),
        ("OWASP", owasp),
        ("Endpoint", endpoint),
        ("Auth Required", auth),
    ]
    if records:
        data.append(("Records Exposed", records))

    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (key, val) in enumerate(data):
        cell_key = table.rows[i].cells[0]
        cell_val = table.rows[i].cells[1]
        cell_key.text = key
        cell_val.text = val

        # Bold keys
        for p in cell_key.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        for p in cell_val.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

        # Color severity
        if key == "Severity":
            colors = {"CRITICAL": "DC2626", "HIGH": "EA580C", "MEDIUM": "CA8A04"}
            bg = colors.get(val.upper(), "64748B")
            set_cell_shading(cell_val, bg)
            for p in cell_val.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE
                    r.font.bold = True

        # Header row background
        set_cell_shading(cell_key, "F1F5F9")

    doc.add_paragraph("")


def add_table_from_data(doc, headers, rows):
    """Add a formatted data table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, "1E293B")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True
                r.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if i % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    doc.add_paragraph("")


def add_callout(doc, text, style="warning"):
    """Add a colored callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    prefix = {"warning": "⚠ ", "critical": "🔴 ", "info": "ℹ "}.get(style, "")
    run = p.add_run(f"{prefix}{text}")
    run.font.size = Pt(10)
    color = {"warning": ORANGE, "critical": RED, "info": BLUE}.get(style, GRAY)
    run.font.color.rgb = color
    run.font.italic = True


# ============================================================================
# DOCUMENT 1: Full Security Assessment Report
# ============================================================================
def build_report():
    doc = Document()
    style_doc(doc)

    add_cover_page(
        doc,
        "MedBasket API\nSecurity Assessment",
        "Comprehensive Vulnerability Report & Remediation Guide",
    )

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Composite Risk Rating: 10.0 / 10.0 — CRITICAL")
    run.font.color.rgb = RED
    run.font.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        "The MedBasket production API (api.medbasket.com) has catastrophic security failures. "
        "Multiple endpoints are publicly accessible without any authentication, exposing the "
        "personal data of 188,000+ users including names, emails, phone numbers, password "
        "reset tokens, OTPs, dates of birth, gender, and physical attributes."
    )
    doc.add_paragraph(
        "A complete attack chain was demonstrated — from zero access to full admin-level "
        "database access in under 3 seconds, fully automated. The chain was proven end-to-end "
        "on the live production system."
    )

    # Scope
    doc.add_heading("Assessment Scope", level=2)
    add_table_from_data(doc,
        ["Component", "Target", "Status"],
        [
            ["API", "api.medbasket.com", "300+ endpoints assessed"],
            ["Admin Panel", "admin.medbasket.com", "Architecture analyzed"],
            ["Storage", "techpepo-development-s3.s3.amazonaws.com", "98,390 files accessible"],
            ["CDN", "media.medbasket.com", "CloudFront distribution"],
        ]
    )

    # Infrastructure
    doc.add_heading("Infrastructure Fingerprint", level=2)
    infra_data = [
        ["Server", "nginx/1.26.2"],
        ["Framework", "FeathersJS (Node.js)"],
        ["Database", "MongoDB (Mongoose ODM)"],
        ["Cloud", "AWS ap-south-1"],
        ["Internal IP", "172.31.XX.XX:3030"],
        ["S3 Bucket", "techpepo-development-s3 (dev bucket in production)"],
        ["Payment", "PayU + Razorpay"],
        ["Auth Strategies", "otp, jwt (local disabled)"],
        ["JWT Algorithm", "HS256, 30-day lifetime"],
        ["JWT Audience", "https://yourdomain.com (FeathersJS default)"],
    ]
    add_table_from_data(doc, ["Component", "Detail"], infra_data)

    # Kill Chain
    doc.add_heading("Proven Kill Chain", level=1)
    doc.add_paragraph(
        "The following attack chain was executed and verified end-to-end during this assessment:"
    )
    chain_steps = [
        ["1", "Enumerate users", "GET /users?$limit=100", "No auth", "188,000+ users with PII"],
        ["2", "Trigger OTP", "POST /users/request-otp", "No auth", "SMS sent to victim"],
        ["3", "Read OTP", "GET /users?$select[]=phoneOtp", "No auth", "OTP code in plaintext"],
        ["4", "Authenticate", "POST /authentication", "No auth", "30-day JWT issued"],
        ["5", "Access orders", "GET /super-admin-users/orders", "Any JWT", "63,000+ orders + GPS"],
        ["6", "Access dashboard", "GET /dashboard", "Any JWT", "Live revenue + customer PII"],
        ["7", "Access medical", "GET /medicine-requests", "Any JWT", "6,655 medical records"],
    ]
    add_table_from_data(doc, ["Step", "Action", "Endpoint", "Auth", "Result"], chain_steps)

    add_callout(doc, "Time from zero access to full admin data: < 3 seconds, fully automatable", "critical")

    # Findings Summary
    doc.add_heading("Findings Summary", level=1)

    findings = [
        ["F-01", "Unauthenticated User Database Access", "CRITICAL", "9.8", "API1 + API2"],
        ["F-02", "Password Hash Exposure via PATCH", "CRITICAL", "9.8", "API3"],
        ["F-03", "Account Takeover — Regular User (OTP)", "CRITICAL", "10.0", "API2"],
        ["F-04", "Account Takeover — Admin User (OTP)", "CRITICAL", "10.0", "API2"],
        ["F-05", "Broken Function-Level Auth (Orders)", "CRITICAL", "9.8", "API5"],
        ["F-06", "Admin Dashboard Without RBAC", "CRITICAL", "9.8", "API5"],
        ["F-07", "Medical Data Exposure", "HIGH", "7.5", "API1"],
        ["F-08", "Swagger/OpenAPI Exposure", "HIGH", "7.5", "API8"],
        ["F-09", "Unauthenticated Write Operations", "HIGH", "8.6", "API2"],
        ["F-10", "Payment Webhook Abuse", "HIGH", "8.1", "API8"],
        ["F-11", "S3 Bucket Misconfiguration", "HIGH", "7.5", "API8"],
        ["F-12", "Schema Disclosure via Validation", "MEDIUM", "5.3", "API3"],
        ["F-13", "JWT Configuration Weaknesses", "MEDIUM", "7.1", "API2"],
        ["F-14", "Admin Panel False Security Perimeter", "CRITICAL", "9.8", "API5"],
        ["F-15", "Super-Admin Unauthenticated Endpoints", "MEDIUM", "5.3", "API2"],
    ]
    add_table_from_data(doc, ["ID", "Title", "Severity", "CVSS", "OWASP"], findings)

    # Detailed findings
    doc.add_heading("Detailed Findings", level=1)

    # F-01
    doc.add_heading("F-01: Unauthenticated User Database Access", level=2)
    add_finding_table(doc, "F-01", "Unauthenticated User Database Access",
        "CRITICAL", "9.8 (AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N)",
        "API1:2023 BOLA + API2:2023 Broken Authentication",
        "GET /users", "None", "188,147 users")
    doc.add_paragraph(
        "The /users endpoint returns all user records without any authentication. "
        "Records include email, phone number, date of birth, gender, health data (height/weight), "
        "live OTP codes, password reset tokens, and account metadata."
    )
    doc.add_paragraph("Evidence — Request:")
    add_code_block(doc, 'GET /users?$limit=3&$select[]=email&$select[]=phoneOtp&$select[]=passwordResetToken\nHost: api.medbasket.com')
    doc.add_paragraph(
        "Users are searchable by any field: email, phone number, or any database field. "
        "No rate limiting is enforced. Pagination via $limit and $skip allows full extraction."
    )

    exposed_fields = [
        ["email", "PII", "User email address"],
        ["phoneNumber", "PII", "Phone number"],
        ["phoneOtp", "CRITICAL", "Live SMS OTP code — enables account takeover"],
        ["passwordResetToken", "CRITICAL", "Active reset token — enables password reset"],
        ["dateOfBirth", "PII", "Date of birth"],
        ["gender", "PII", "Gender"],
        ["height / weight", "Health", "Physical health data"],
        ["rewardCoinsBalance", "Financial", "In-app currency"],
    ]
    add_table_from_data(doc, ["Field", "Sensitivity", "Description"], exposed_fields)

    # F-02
    doc.add_heading("F-02: Password Hash Exposure via PATCH", level=2)
    add_finding_table(doc, "F-02", "Password Hash Exposure via PATCH",
        "CRITICAL", "9.8 (AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N)",
        "API3:2023 Broken Object Property Level Authorization",
        "PATCH /users/{id}", "None",
        "All 188,147 users' bcrypt hashes extractable")
    doc.add_paragraph(
        "Sending an empty PATCH request ({}) to any user ID returns the raw Mongoose document, "
        "bypassing all FeathersJS protect() hooks. The response includes the bcrypt password hash, "
        "internal Mongoose metadata ($__, $isNew, _doc), and all fields."
    )
    doc.add_paragraph(
        "This also performs a database WRITE — the updatedAt timestamp changes, confirming "
        "unauthenticated database modification."
    )
    doc.add_paragraph("Confirmed on multiple accounts:")
    add_table_from_data(doc, ["User ID", "Email", "Hash Prefix"],
        [
            ["600000a1...", "admin@[REDACTED].com", "$2a$10$[HASH]..."],
            ["600000a5...", "testuser2@[REDACTED].com", "$2a$10$[HASH]..."],
            ["600000a4...", "testuser@[REDACTED].com", "$2a$10$[HASH]..."],
        ]
    )

    # F-03
    doc.add_heading("F-03: Account Takeover — Regular User (OTP Chain)", level=2)
    add_finding_table(doc, "F-03", "Account Takeover — Regular User",
        "CRITICAL", "10.0 (AV:N/AC:L/PR:N/UI:N/S:C/C:H/I:H/A:H)",
        "API2:2023 Broken Authentication",
        "Multiple endpoints", "None",
        "Any of 188,147 accounts compromisable")

    doc.add_paragraph("The complete account takeover chain in 4 API calls:")
    steps = [
        ["Step 1", "GET /users?email=victim@example.com", "Get target's phone number"],
        ["Step 2", "POST /users/request-otp", "Trigger OTP SMS to victim's phone"],
        ["Step 3", "GET /users?$select[]=phoneOtp", "Read OTP code from API (no auth)"],
        ["Step 4", "POST /authentication", "Authenticate with stolen OTP → 30-day JWT"],
    ]
    add_table_from_data(doc, ["Step", "Request", "Result"], steps)
    add_callout(doc, "Time to compromise: < 3 seconds. Works for ANY user. Fully automatable.", "critical")
    doc.add_paragraph(
        "The authentication response also leaks the user's bcrypt password hash in the response body."
    )

    # F-04
    doc.add_heading("F-04: Account Takeover — Admin User", level=2)
    add_finding_table(doc, "F-04", "Account Takeover — Admin User",
        "CRITICAL", "10.0 (AV:N/AC:L/PR:N/UI:N/S:C/C:H/I:H/A:H)",
        "API2:2023 Broken Authentication",
        "Multiple endpoints", "None",
        "Admin account (admin@[REDACTED].com)")
    doc.add_paragraph(
        "The same OTP chain was executed against the confirmed super-admin account. "
        "Both regular and admin JWTs return identical data from admin endpoints — "
        "confirming zero role-based access control."
    )

    # F-05
    doc.add_heading("F-05: Broken Function-Level Authorization (Orders)", level=2)
    add_finding_table(doc, "F-05", "Broken Function-Level Authorization",
        "CRITICAL", "9.8 (AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N)",
        "API5:2023 Broken Function Level Authorization",
        "GET /super-admin-users/orders", "Any valid JWT (no role check)",
        "62,900+ orders")
    doc.add_paragraph(
        "Despite the URL path suggesting admin-only access, any authenticated JWT "
        "returns the full order dataset. Each order contains:"
    )
    pii_list = [
        "Customer name and phone number",
        "Full delivery address (street, city, state, postal code)",
        "GPS coordinates (latitude, longitude)",
        "Alternate phone number",
        "Payment details, coupon codes, item list",
    ]
    for item in pii_list:
        p = doc.add_paragraph(item, style="List Bullet")

    # F-06
    doc.add_heading("F-06: Admin Dashboard Without RBAC", level=2)
    add_finding_table(doc, "F-06", "Admin Dashboard Without RBAC",
        "CRITICAL", "9.8", "API5:2023 Broken Function Level Authorization",
        "GET /dashboard", "Any valid JWT (no role check)")
    doc.add_paragraph(
        "The admin dashboard returns live business intelligence: monthly revenue (INR 18.6L), "
        "order counts (3,335), customer counts (16,140), weekly breakdowns, top-selling products, "
        "and recent orders with real customer names and email addresses."
    )

    # F-07
    doc.add_heading("F-07: Medical Data Exposure", level=2)
    add_finding_table(doc, "F-07", "Medical Data Exposure",
        "HIGH", "7.5", "API1:2023 BOLA",
        "GET /medicine-requests", "Any valid JWT",
        "6,655 medicine requests")
    doc.add_paragraph(
        "Medicine requests with prescription file URLs are accessible to any authenticated user. "
        "This constitutes protected health information under the DPDP Act with the highest "
        "protection requirements."
    )

    # F-08 through F-15 (condensed)
    doc.add_heading("F-08: Swagger/OpenAPI Exposure", level=2)
    doc.add_paragraph(
        "The complete API specification is publicly accessible at /swagger.json — 300+ endpoints "
        "documented with schemas, auth requirements, and the internal AWS server URL "
        "(http://ip-172-31-XX-XX.ap-south-1.compute.internal:3030)."
    )

    doc.add_heading("F-09: Unauthenticated Write Operations", level=2)
    writes = [
        ["PATCH /users/{id}", "200 OK", "DB write + hash leak"],
        ["POST /admin-zip-codes", "201 Created", "Admin data injection"],
        ["POST /users/reset-password", "201 Created", "Reset record creation"],
        ["POST /users/request-otp", "201 Created", "OTP generation + SMS send"],
        ["POST /request-super-admin-otp", "201 Created", "Admin OTP generation"],
    ]
    add_table_from_data(doc, ["Endpoint", "Response", "Impact"], writes)

    doc.add_heading("F-10: Payment Webhook Abuse", level=2)
    doc.add_paragraph(
        "Payment webhooks accept unverified payloads. POST /webhooks returns pricing data (200 OK). "
        "POST /webhooks/delivery-tracking accepts fake delivery updates. Error messages from "
        "PayU and Razorpay webhooks disclose integration details."
    )

    doc.add_heading("F-11: S3 Bucket Misconfiguration", level=2)
    doc.add_paragraph(
        "The S3 bucket 'techpepo-development-s3' (a development bucket used in production) "
        "has publicly readable objects. 98,390 file URLs are enumerable via /attachments "
        "without authentication. AWS Access Key ID is exposed in pre-signed URLs."
    )

    doc.add_heading("F-12: Schema Disclosure via Validation Errors", level=2)
    doc.add_paragraph(
        "Invalid $select queries return all 27 allowed field names including undocumented fields: "
        "googleId, profileToken, profileTokenValidTill, identifierType, isWeb, referralCode."
    )

    doc.add_heading("F-13: JWT Configuration Weaknesses", level=2)
    jwt_issues = [
        ["Default audience", "https://yourdomain.com — FeathersJS default never changed"],
        ["30-day token lifetime", "Industry standard: 15-60 minutes"],
        ["No refresh tokens", "Only access tokens issued"],
        ["Password in auth response", "Bcrypt hash returned on successful login"],
        ["HS256 algorithm", "Symmetric key — single secret = all tokens forgeable"],
        ["No revocation", "No blacklist mechanism"],
    ]
    add_table_from_data(doc, ["Issue", "Detail"], jwt_issues)

    doc.add_heading("F-14: Admin Panel False Security Perimeter", level=2)
    doc.add_paragraph(
        "The admin panel (admin.medbasket.com) uses Auth.js v5 with a separate super-admin-users "
        "collection. The frontend authentication is properly implemented. However, ALL admin API "
        "endpoints are accessible with any regular user JWT — the admin panel protects the UI, "
        "not the data. An attacker bypasses the admin panel entirely by calling the API directly."
    )

    doc.add_heading("F-15: Super-Admin Unauthenticated Endpoints", level=2)
    sa_endpoints = [
        ["POST /super-admin-users/forgot-password", "No", "200 — sends reset email, confirms admin exists"],
        ["POST /super-admin-users/reset-password", "No", "400 — validates schema, discloses field names"],
        ["POST /request-super-admin-otp", "No", "Requires phoneNumber — can trigger admin OTP"],
    ]
    add_table_from_data(doc, ["Endpoint", "Auth Required", "Response"], sa_endpoints)

    # Data Exposure Summary
    doc.add_heading("Total Data Exposure", level=1)
    exposure = [
        ["User accounts (with PII)", "188,147", "CRITICAL"],
        ["Bcrypt password hashes", "188,147", "CRITICAL"],
        ["Customer orders (with addresses + GPS)", "62,900+", "CRITICAL"],
        ["Medicine requests (health data)", "6,655", "HIGH"],
        ["S3 file attachments", "98,390", "HIGH"],
        ["API endpoints documented", "300+", "HIGH"],
    ]
    add_table_from_data(doc, ["Dataset", "Records", "Severity"], exposure)

    add_callout(doc,
        "All data is extractable in full using simple pagination scripts. "
        "No rate limiting, no anomaly detection, no alerts. "
        "An attacker can download the entire database in under 90 minutes.",
        "critical"
    )

    # Regulatory
    doc.add_heading("Regulatory Impact", level=1)

    doc.add_heading("India DPDP Act 2023", level=2)
    reg_items = [
        ["Section 8(5)", "Failure to implement reasonable security safeguards"],
        ["Section 12", "Mandatory breach notification to Data Protection Board"],
        ["Penalty", "Up to INR 250 crore (~$30M USD)"],
    ]
    add_table_from_data(doc, ["Section", "Violation"], reg_items)

    doc.add_heading("GDPR (if EU users exist)", level=2)
    gdpr = [
        ["Article 5(1)(f)", "Violation of integrity and confidentiality principle"],
        ["Article 32", "Failure to implement appropriate technical measures"],
        ["Article 33", "Mandatory 72-hour breach notification"],
        ["Penalty", "Up to 4% of annual global turnover or EUR 20 million"],
    ]
    add_table_from_data(doc, ["Article", "Violation"], gdpr)

    # Remediation
    doc.add_heading("Remediation Plan", level=1)

    doc.add_heading("EMERGENCY — Within Hours", level=2)
    emergency = [
        "Take API offline or behind VPN/firewall",
        "Apply authenticate('jwt') hook to ALL FeathersJS services",
        "Fix PATCH response — serialize Mongoose documents, protect() on all methods",
        "Strip sensitive fields from all API responses (phoneOtp, passwordResetToken, password)",
        "Disable Swagger (/swagger.json, /docs) in production",
    ]
    for item in emergency:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 48 Hours", level=2)
    short = [
        "Force password reset for all 188K users",
        "Rotate JWT secret (invalidates all sessions)",
        "Rotate AWS access keys",
        "Add RBAC to admin endpoints (/orders, /dashboard, /medicine-requests)",
        "Block unauthenticated POST on write endpoints",
        "Secure super-admin forgot-password/reset-password endpoints",
        "Webhook signature verification (PayU hash, Razorpay HMAC)",
    ]
    for item in short:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 Week", level=2)
    week = [
        "Implement rate limiting (100 req/min per IP, 5/min on auth)",
        "Add security headers (HSTS, CSP, X-Content-Type-Options)",
        "Remove internal URLs from configs",
        "Migrate from dev S3 bucket to production",
        "Reduce JWT lifetime to 15 minutes + refresh token",
    ]
    for item in week:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 Month", level=2)
    month = [
        "Full security audit of all 300+ endpoints",
        "Implement API gateway (centralized auth, rate limiting, WAF)",
        "Set up monitoring and alerting for bulk data access",
        "Third-party penetration test",
        "DPDP Act compliance review",
        "User notification (mandatory under DPDP Act Section 12)",
    ]
    for item in month:
        doc.add_paragraph(item, style="List Bullet")

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The MedBasket API has no effective security controls on its most sensitive endpoints. "
        "This assessment demonstrated a complete, automated kill chain from zero access to full "
        "admin-level database access — 188,000 users' personal data, 63,000 customer orders "
        "with physical addresses and GPS coordinates, 6,655 medical records, and live business "
        "intelligence. The API must be taken offline immediately."
    )

    p = doc.add_paragraph()
    run = p.add_run("Every user credential, every order, every address, and every medical "
                     "record must be considered compromised.")
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_paragraph("")
    doc.add_paragraph("Full assessment conducted 2026-02-18 as part of authorized security review.")

    # Signatures
    doc.add_heading("Signatures", level=1)
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    sigs = [
        ("Assessment conducted by:", ""),
        ("Date:", "2026-02-18"),
        ("Authorization reference:", ""),
        ("Client representative:", ""),
    ]
    for i, (label, val) in enumerate(sigs):
        sig_table.rows[i].cells[0].text = label
        sig_table.rows[i].cells[1].text = val
        sig_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    path = os.path.join(OUT_DIR, "MedBasket_Security_Assessment_Report.docx")
    doc.save(path)
    print(f"  Report:     {path}")
    return path


# ============================================================================
# DOCUMENT 2: Client Walkthrough
# ============================================================================
def build_walkthrough():
    doc = Document()
    style_doc(doc)

    add_cover_page(
        doc,
        "MedBasket API\nClient Walkthrough",
        "Presentation Guide for Security Assessment Findings",
    )

    doc.add_heading("Meeting Agenda (75-90 minutes)", level=1)
    agenda = [
        ["5 min", "Introduction", "Scope, methodology, authorization review"],
        ["10 min", "Phase 1", "Unauthenticated data exposure — 188K users (live demo)"],
        ["10 min", "Phase 2", "Account takeover via OTP chain (live demo)"],
        ["15 min", "Phase 3", "Privilege escalation — regular user → admin data (live demo)"],
        ["5 min", "Phase 4", "Additional attack surface — webhooks, S3, writes"],
        ["10 min", "Phase 5", "Admin panel false perimeter analysis"],
        ["10 min", "Phase 6", "Full data extraction + visual dashboard demo"],
        ["10 min", "Remediation", "Priority fixes and timeline"],
        ["5 min", "Q&A", ""],
    ]
    add_table_from_data(doc, ["Time", "Phase", "Description"], agenda)

    # Pre-meeting
    doc.add_heading("Pre-Meeting Setup", level=1)
    doc.add_heading("Requirements", level=2)
    reqs = [
        "Terminal with curl and python3",
        "Browser with poc_dashboard.html open (pre-loaded, no internet needed)",
        "Screen sharing enabled",
        "Recording enabled (for client's records)",
        "Written authorization document on screen",
    ]
    for r in reqs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Pre-Meeting Commands", level=2)
    add_code_block(doc,
        "# Generate fresh dashboard (15 min before meeting)\n"
        "python3 build_dashboard.py\n"
        "# Open in browser\n"
        "open poc_dashboard.html\n\n"
        "# Run interactive PoC\n"
        "chmod +x poc_medbasket.sh\n"
        "./poc_medbasket.sh"
    )

    # Opening
    doc.add_heading("Opening Statement", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '"Thank you for joining. Today we\'ll walk through critical security findings '
        'in the MedBasket production API at api.medbasket.com. We identified 15 vulnerabilities '
        '— 7 Critical, 5 High, 2 Medium, 1 Informational. What I\'m about to show you '
        'represents an active, exploitable data breach affecting over 188,000 users, '
        '63,000 customer orders, and 6,600 medical records. Every vulnerability was '
        'verified on your live production system. The composite CVSS risk rating is '
        '10.0 out of 10.0."'
    )
    run.font.italic = True
    run.font.color.rgb = DARK_BLUE

    # Phase 1
    doc.add_heading("Phase 1: Unauthenticated Data Exposure", level=1)
    doc.add_paragraph("Findings covered: F-01, F-02, F-08, F-12 | Time: 10 minutes")

    doc.add_heading("Demo 1.1 — User Database [F-01]", level=2)
    doc.add_paragraph("Show in terminal:")
    add_code_block(doc, 'curl -s "https://api.medbasket.com/users?$limit=1" | python3 -m json.tool | head -30')
    doc.add_heading("Talking Points:", level=3)
    points = [
        "\"This is a simple GET request. No API key, no login, no token.\"",
        "Point out the total field: \"188,000+ users are accessible this way.\"",
        "Point out passwordResetToken: \"This is a live password reset token.\"",
        "Point out phoneOtp: \"This is the live SMS verification code.\"",
        "\"An attacker pages through ALL users by changing the $skip parameter.\"",
    ]
    for pt in points:
        doc.add_paragraph(pt, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run('Impact: "Anyone on the internet can download your entire user database right now."')
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_heading("Demo 1.2 — Swagger [F-08]", level=2)
    add_code_block(doc, 'curl -s "https://api.medbasket.com/swagger.json" | python3 -c "\nimport json,sys\nd = json.load(sys.stdin)\nprint(f\'Endpoints: {len(d[\"paths\"])}\')\nprint(f\'Internal server: {d[\"servers\"][0][\"url\"]}\')\n"')
    doc.add_paragraph("300+ endpoints documented publicly. Internal AWS IP 172.31.XX.XX disclosed.")

    doc.add_heading("Demo 1.3 — PATCH Hash Leak [F-02]", level=2)
    doc.add_paragraph("Show cached output (this is a write operation):")
    add_code_block(doc, 'curl -s -X PATCH "https://api.medbasket.com/users/<USER_ID>" \\\n  -H "Content-Type: application/json" -d \'{}\' | python3 -m json.tool | head -20')
    doc.add_paragraph(
        "Sending an empty PATCH to any user returns their full Mongoose document "
        "including the bcrypt password hash. FeathersJS protect() hooks bypassed on PATCH."
    )

    # Phase 2
    doc.add_heading("Phase 2: Account Takeover — OTP Chain", level=1)
    doc.add_paragraph("Findings covered: F-03, F-04 | Time: 10 minutes")

    doc.add_heading("Demo 2.1 — Complete Chain [F-03]", level=2)
    doc.add_paragraph("Walk through each step in terminal:")
    add_code_block(doc,
        "# Step 1: Find target\n"
        "curl -s \"https://api.medbasket.com/users?$limit=1&$select[]=phoneNumber\"\n\n"
        "# Step 2: Request OTP (no auth)\n"
        "curl -s -X POST \"https://api.medbasket.com/users/request-otp\" \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"phoneNumber\":\"+91XXXXXXXXXX\"}'\n\n"
        "# Step 3: Read OTP (no auth) — THE CRITICAL STEP\n"
        "curl -s \"https://api.medbasket.com/users?phoneNumber=%2B91XXXXXXXXXX&$select[]=phoneOtp\"\n\n"
        "# Step 4: Authenticate with stolen OTP\n"
        "curl -s -X POST \"https://api.medbasket.com/authentication\" \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"strategy\":\"otp\",\"phoneNumber\":\"+91XXXXXXXXXX\",\"otp\":\"XXXXX\"}'"
    )
    doc.add_heading("Step-by-Step Talking Points:", level=3)
    chain_points = [
        "\"First, I pick any user from the database. I have their phone number.\"",
        "\"Now I request an OTP for their phone. This sends a real SMS.\"",
        "\"Here's the critical part — I read the OTP directly from your API. "
        "The same API that sent the SMS shows the code to anyone.\"",
        "\"I authenticate with the stolen OTP. JWT valid for 30 days.\"",
        "\"The auth response also leaks the user's bcrypt password hash.\"",
    ]
    for pt in chain_points:
        doc.add_paragraph(pt, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run(
        'Impact: "I can take over any account in under 3 seconds, fully automated. '
        'This works for all 188,000 users."'
    )
    run.font.bold = True
    run.font.color.rgb = RED

    # Phase 3
    doc.add_heading("Phase 3: Privilege Escalation", level=1)
    doc.add_paragraph("Findings covered: F-05, F-06, F-07, F-13 | Time: 15 minutes")

    doc.add_heading("Demo 3.1 — Admin Dashboard [F-06]", level=2)
    add_code_block(doc, 'TOKEN="<jwt_from_step_4>"\ncurl -s "https://api.medbasket.com/dashboard" \\\n  -H "Authorization: Bearer $TOKEN" | python3 -m json.tool')
    doc.add_paragraph(
        "\"Using the JWT from a regular user account — not an admin. "
        "Live monthly revenue, order counts, customer metrics, and recent orders "
        "with real customer names and email addresses.\""
    )

    doc.add_heading("Demo 3.2 — All Orders [F-05]", level=2)
    add_code_block(doc, 'curl -s "https://api.medbasket.com/super-admin-users/orders?$limit=1" \\\n  -H "Authorization: Bearer $TOKEN" | python3 -m json.tool | head -50')
    doc.add_paragraph(
        "\"Despite the URL '/super-admin-users/', there is no role check. "
        "62,900+ orders accessible with full delivery addresses and GPS coordinates.\""
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Impact: "Anyone who creates an account can see every order your platform '
        'has ever processed, including where your customers live."'
    )
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_heading("Demo 3.3 — Medicine Requests [F-07]", level=2)
    doc.add_paragraph("6,655 medicine requests with prescription file URLs. Protected health information.")

    # Phase 4
    doc.add_heading("Phase 4: Additional Attack Surface", level=1)
    doc.add_paragraph("Findings covered: F-09, F-10, F-11 | Time: 5 minutes")
    add_table_from_data(doc,
        ["Finding", "Evidence", "Severity"],
        [
            ["F-09: Unauthenticated writes", "POST /admin-zip-codes = 201, PATCH = DB write", "HIGH"],
            ["F-10: Webhook abuse", "No signature verification, error message leaks", "HIGH"],
            ["F-11: S3 misconfiguration", "98K files accessible, dev bucket in prod", "HIGH"],
            [".env exists", "GET /.env = 403 (not 404)", "INFO"],
            ["Server version", "nginx/1.26.2 in headers", "LOW"],
        ]
    )

    # Phase 5
    doc.add_heading("Phase 5: Admin Panel — False Perimeter", level=1)
    doc.add_paragraph("Findings covered: F-14, F-15 | Time: 10 minutes")
    doc.add_paragraph(
        "\"The admin panel at admin.medbasket.com uses Auth.js v5 — proper, modern authentication. "
        "The super-admin-users collection is separate. The login screen is well implemented. "
        "BUT — it doesn't matter. Every admin API endpoint works with a regular user JWT. "
        "The admin panel's security is a false perimeter — it protects the UI, not the data.\""
    )

    # Phase 6
    doc.add_heading("Phase 6: Full Data Extraction + Dashboard", level=1)
    doc.add_paragraph("Time: 10 minutes")

    doc.add_heading("Extraction Scripts", level=2)
    add_table_from_data(doc,
        ["Script", "Records", "Output"],
        [
            ["fetch_orders.py", "~63,000 orders", "data/orders.json + .csv"],
            ["fetch_users.py", "~188,000 users", "data/users.json + .csv"],
            ["fetch_medicine_requests.py", "~6,655 records", "data/medicine_requests.json + .csv"],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'Impact: "An attacker can extract your entire database in under 90 minutes. '
        'No rate limits, no monitoring, no alerts. Clean CSV files ready for the dark web."'
    )
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_heading("Visual Dashboard Demo", level=2)
    doc.add_paragraph(
        "Open poc_dashboard.html in the browser. Show 500 orders with expandable PII, "
        "500 users with leaked OTPs, 200 medicine requests. Point out the red totals: "
        "63,000+ orders, 188,000+ users, 6,655 medical records."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        '"This is what an attacker\'s dashboard looks like. It took 30 seconds to generate '
        'from your live production data."'
    )
    run.font.italic = True
    run.font.color.rgb = DARK_BLUE

    # Remediation
    doc.add_heading("Remediation Timeline", level=1)

    doc.add_heading("EMERGENCY — Within Hours", level=2)
    for item in [
        "Take API offline or behind VPN/firewall",
        "Apply authenticate('jwt') hook to ALL services",
        "Fix PATCH response serialization",
        "Strip sensitive fields from all API responses",
        "Disable Swagger in production",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 48 Hours", level=2)
    for item in [
        "Force password reset for all 188K users",
        "Rotate JWT secret + AWS access keys",
        "Add RBAC to admin endpoints",
        "Secure super-admin endpoints",
        "Webhook signature verification",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 Week", level=2)
    for item in [
        "Rate limiting, security headers, JWT lifetime reduction",
        "Remove internal URLs, migrate S3 bucket",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 Month", level=2)
    for item in [
        "Full 300+ endpoint audit, API gateway, monitoring",
        "Third-party pentest, DPDP compliance, user notification",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Closing
    doc.add_heading("Closing Statement", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '"To summarize: your API has no effective security boundary. We demonstrated a complete '
        'chain from zero access to full admin data in 3 seconds. We built scripts that extract '
        'your entire database — 188,000 users, 63,000 orders with GPS, and 6,600 medical records. '
        'Your admin panel protects only the UI. The most urgent action is to take the API offline '
        'until authentication is enforced on every endpoint. Every hour increases your regulatory '
        'exposure and your users\' risk."'
    )
    run.font.italic = True
    run.font.color.rgb = DARK_BLUE

    path = os.path.join(OUT_DIR, "MedBasket_PoC_Walkthrough.docx")
    doc.save(path)
    print(f"  Walkthrough: {path}")
    return path


# ============================================================================
# DOCUMENT 3: Findings Evidence Package
# ============================================================================
def build_evidence():
    doc = Document()
    style_doc(doc)

    add_cover_page(
        doc,
        "MedBasket API\nFindings Evidence Package",
        "Raw Evidence: JWT Tokens, API Responses, Access Matrices",
    )

    doc.add_heading("Document Purpose", level=1)
    doc.add_paragraph(
        "This evidence package contains verified findings, JWT tokens, and raw API responses "
        "captured during the security assessment of api.medbasket.com. All evidence was captured "
        "on 2026-02-18 between 05:55 UTC and 06:35 UTC."
    )

    # Infrastructure
    doc.add_heading("Infrastructure Fingerprint", level=1)
    infra = [
        ["Target", "api.medbasket.com"],
        ["Server", "nginx/1.26.2"],
        ["Framework", "FeathersJS (Node.js)"],
        ["Database", "MongoDB (Mongoose ODM)"],
        ["Cloud", "AWS ap-south-1"],
        ["Internal IP", "172.31.XX.XX:3030"],
        ["S3 Bucket", "techpepo-development-s3"],
        ["CDN", "media.medbasket.com (CloudFront)"],
        ["Payment", "PayU, Razorpay"],
        ["Auth", "otp, jwt (local disabled)"],
        ["JWT", "HS256, 30-day, aud: https://yourdomain.com"],
    ]
    add_table_from_data(doc, ["Component", "Detail"], infra)

    # F-01
    doc.add_heading("F-01: Unauthenticated User Database Access", level=1)
    add_finding_table(doc, "F-01", "User Database Access", "CRITICAL",
        "9.8", "API1 + API2", "GET /users", "None", "188,147 users")
    doc.add_paragraph("Request:")
    add_code_block(doc, "GET /users?$limit=3&$select[]=_id&$select[]=email&$select[]=phoneOtp&$select[]=passwordResetToken\nHost: api.medbasket.com\nAccept: application/json")
    doc.add_paragraph("Response (captured 2026-02-18 ~05:57 UTC):")
    add_code_block(doc, '{\n  "total": 188147,\n  "data": [\n    {\n      "_id": "6000000000000000000000a1",\n      "email": "admin@[REDACTED].com",\n      "phoneOtp": null,\n      "passwordResetToken": "[RESET_TOKEN_REDACTED]"\n    },\n    {\n      "_id": "6000000000000000000000a2",\n      "email": "user1@[REDACTED].com",\n      "phoneOtp": "XXXXX"\n    }\n  ]\n}')

    # F-02
    doc.add_heading("F-02: Password Hash Exposure via PATCH", level=1)
    add_finding_table(doc, "F-02", "PATCH Hash Leak", "CRITICAL",
        "9.8", "API3", "PATCH /users/{id}", "None")
    doc.add_paragraph("Request:")
    add_code_block(doc, "PATCH /users/6000000000000000000000a1\nHost: api.medbasket.com\nContent-Type: application/json\n\n{}")
    doc.add_paragraph("Response includes bcrypt hash in _doc.password:")
    add_code_block(doc, '"_doc": {\n  "email": "admin@[REDACTED].com",\n  "password": "$2a$10$[HASH_REDACTED]",\n  "phoneOtp": null,\n  "passwordResetToken": "[RESET_TOKEN]..."\n}')
    doc.add_paragraph("Confirmed hashes:")
    add_table_from_data(doc, ["User ID", "Email", "Hash Prefix"],
        [
            ["600000a1...", "admin@[REDACTED].com", "$2a$10$[HASH]..."],
            ["600000a5...", "testuser2@[REDACTED].com", "$2a$10$[HASH]..."],
            ["600000a4...", "testuser@[REDACTED].com", "$2a$10$[HASH]..."],
        ])

    # F-03
    doc.add_heading("F-03: Account Takeover — Regular User", level=1)
    add_finding_table(doc, "F-03", "OTP Chain Takeover", "CRITICAL",
        "10.0", "API2", "Multiple", "None", "Any of 188,147 users")
    doc.add_paragraph("Step 1 — Request OTP:")
    add_code_block(doc, 'POST /users/request-otp\n{"phoneNumber":"+91XXXXXXXXXX"}\n\n→ 201 Created: {"message":"OTP has been sent to your phone number."}')
    doc.add_paragraph("Step 2 — Read OTP:")
    add_code_block(doc, 'GET /users?phoneNumber=%2B91XXXXXXXXXX&$select[]=phoneOtp\n\n→ {"data": [{"phoneOtp": "XXXXX"}]}')
    doc.add_paragraph("Step 3 — Authenticate:")
    add_code_block(doc, 'POST /authentication\n{"strategy":"otp","phoneNumber":"+91XXXXXXXXXX","otp":"XXXXX"}\n\n→ 201 Created, JWT issued + password hash leaked in response')

    doc.add_heading("JWT Token — Regular User", level=2)
    add_code_block(doc, "<TOKEN_REDACTED>")
    add_table_from_data(doc, ["Field", "Value"],
        [
            ["Subject (user ID)", "6000000000000000000000a4"],
            ["Issued At", "2026-02-18 06:16:27 UTC"],
            ["Expires", "2026-03-20 06:16:27 UTC"],
            ["Audience", "https://yourdomain.com (default)"],
        ])

    # F-04
    doc.add_heading("F-04: Account Takeover — Admin User", level=1)
    add_finding_table(doc, "F-04", "Admin OTP Takeover", "CRITICAL",
        "10.0", "API2", "Multiple", "None", "admin@[REDACTED].com")
    doc.add_paragraph("Same OTP chain against admin phone +91XXXXXXXXXX. OTP captured: XXXXX.")
    doc.add_heading("JWT Token — Admin User", level=2)
    add_code_block(doc, "<TOKEN_REDACTED>")
    doc.add_paragraph("Both regular and admin JWTs return identical data from /super-admin-users/orders — zero RBAC.")

    # F-05
    doc.add_heading("F-05: Orders — No RBAC", level=1)
    add_finding_table(doc, "F-05", "Broken Function-Level Auth", "CRITICAL",
        "9.8", "API5", "GET /super-admin-users/orders", "Any JWT", "62,900+ orders")
    doc.add_paragraph("Sample order data (PII redacted for evidence):")
    add_code_block(doc, '{\n  "total": 62880,\n  "data": [{\n    "orderId": "0XXXXXXXX",\n    "status": "paid",\n    "paymentAmount": 729,\n    "address": {\n      "userName": "S***** S*****",\n      "phoneNumber": "77971*****",\n      "city": "[REDACTED_DISTRICT]",\n      "coordinates": {"latitude": XX.XXXXXX, "longitude": XX.XXXXXX}\n    }\n  }]\n}')

    # F-06, F-07
    doc.add_heading("F-06: Dashboard Without RBAC", level=1)
    doc.add_paragraph("GET /dashboard with regular user JWT returns: Revenue INR 18,63,460.81 (current), "
                       "orders 3,335 (current), customers 16,140, recent orders with customer PII.")

    doc.add_heading("F-07: Medicine Requests", level=1)
    doc.add_paragraph("GET /medicine-requests returns 6,655 records with prescription file URLs.")

    # F-08 through F-15
    doc.add_heading("F-08 to F-13: Additional Findings", level=1)
    additional = [
        ["F-08", "Swagger exposed", "300+ endpoints, internal IP"],
        ["F-09", "Unauthenticated writes", "5 endpoints accept writes without auth"],
        ["F-10", "Webhook abuse", "No signature verification, error leaks"],
        ["F-11", "S3 misconfiguration", "98,390 files, dev bucket, public objects"],
        ["F-12", "Schema disclosure", "27 fields disclosed via validation errors"],
        ["F-13", "JWT weaknesses", "Default config, 30-day lifetime, HS256"],
    ]
    add_table_from_data(doc, ["ID", "Finding", "Detail"], additional)

    doc.add_heading("F-14: Admin Panel False Perimeter", level=1)
    doc.add_paragraph(
        "admin.medbasket.com uses Auth.js v5 with separate super-admin-users collection. "
        "Frontend properly secured. However, all admin API data accessible with regular user JWT. "
        "Evidence: __Host-authjs.csrf-token cookie, 404 on /super-admin-users with regular JWT, "
        "but 200 on /super-admin-users/orders with same JWT."
    )

    doc.add_heading("F-15: Super-Admin Unauthenticated Endpoints", level=1)
    sa = [
        ["POST /super-admin-users/forgot-password", "No", "200 — \"Sent reset password mail\""],
        ["POST /super-admin-users/reset-password", "No", "400 — schema validation only"],
        ["POST /request-super-admin-otp", "No", "Requires phoneNumber"],
    ]
    add_table_from_data(doc, ["Endpoint", "Auth", "Response"], sa)

    # Endpoint matrix
    doc.add_heading("Complete Endpoint Access Matrix", level=1)

    doc.add_heading("No Authentication Required", level=2)
    no_auth = [
        ["/users", "GET", "188,147", "User PII, OTPs, reset tokens"],
        ["/users/{id}", "PATCH", "per-user", "Password hashes, Mongoose internals"],
        ["/attachments", "GET", "98,390", "S3 file URLs"],
        ["/admin-zip-codes", "GET/POST", "19,794", "Delivery zones + GPS"],
        ["/swagger.json", "GET", "300+", "Full API specification"],
        ["/products", "GET", "catalog", "Product data, S3 URLs with AWS keys"],
    ]
    add_table_from_data(doc, ["Endpoint", "Method", "Records", "Data Type"], no_auth)

    doc.add_heading("Any Authenticated User (No Role Check)", level=2)
    any_auth = [
        ["/super-admin-users/orders", "GET", "62,880", "Orders + addresses + GPS"],
        ["/dashboard", "GET", "live", "Revenue, customers, recent orders"],
        ["/medicine-requests", "GET", "6,655", "Medical requests + prescriptions"],
        ["/medicine-remainder", "GET", "708", "Medicine reminders"],
        ["/referral-credits", "GET", "102", "Financial data"],
    ]
    add_table_from_data(doc, ["Endpoint", "Method", "Records", "Data Type"], any_auth)

    # Data extraction
    doc.add_heading("Full Data Extraction Evidence", level=1)
    doc.add_paragraph(
        "Standalone scripts were built to demonstrate complete database extraction. "
        "All data paginated via $limit and $skip with no rate limiting encountered."
    )
    extract = [
        ["fetch_orders.py", "~63,000 orders", "data/orders.json + .csv", "Names, phones, addresses, GPS"],
        ["fetch_users.py", "~188,000 users", "data/users.json + .csv", "Emails, OTPs, hashes, DOBs"],
        ["fetch_medicine_requests.py", "~6,655 records", "data/medicine_requests.json + .csv", "Prescriptions, patient IDs"],
    ]
    add_table_from_data(doc, ["Script", "Records", "Output", "PII Exposed"], extract)

    # Signatures
    doc.add_heading("Signatures", level=1)
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Assessment conducted by:", ""),
        ("Date:", "2026-02-18"),
        ("Authorization reference:", ""),
        ("Client representative:", ""),
    ]):
        sig_table.rows[i].cells[0].text = label
        sig_table.rows[i].cells[1].text = val
        sig_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    path = os.path.join(OUT_DIR, "MedBasket_Findings_Evidence.docx")
    doc.save(path)
    print(f"  Evidence:   {path}")
    return path


# ============================================================================
# DOCUMENT 4: Database Schema Map
# ============================================================================
def build_schema_map():
    doc = Document()
    style_doc(doc)

    add_cover_page(
        doc,
        "MedBasket API\nDatabase Schema Map",
        "Complete Schema Reconstruction via API Probing — No Source Code Access",
    )

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Through systematic probing of the MedBasket API (api.medbasket.com), we reconstructed "
        "the majority of the application's database schema without any source code access. The API "
        "is built on FeathersJS + MongoDB (Mongoose) and exposes schema information through multiple channels."
    )

    add_table_from_data(doc,
        ["Technique", "Data Recovered"],
        [
            ["$select[]=__INVALID__ validation", "Field lists for 10+ collections"],
            ["POST empty {} validation chaining", "Required fields for 6+ services"],
            ["GET response analysis", "Full nested schemas for orders, users, stores"],
            ["Mongoose _doc / $__ internal leak", "Raw document internals on /users"],
            ["Swagger/OpenAPI (/swagger.json)", "272 paths, 93 services, 567 methods"],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run("563 of 567 endpoint methods have NO security definition in Swagger. Only 4 methods define any security.")
    run.font.bold = True
    run.font.color.rgb = RED

    # API Surface
    doc.add_heading("API Surface Overview", level=1)
    add_table_from_data(doc,
        ["Metric", "Value"],
        [
            ["Total API paths", "272"],
            ["Total services", "93"],
            ["Total endpoint methods", "567"],
            ["Methods with security defined", "4 (0.7%)"],
            ["Methods without security", "563 (99.3%)"],
            ["Swagger schema definitions", "0 (empty!)"],
        ]
    )

    # Users Schema
    doc.add_heading("Users Collection", level=1)
    doc.add_paragraph("Total: ~188,000 records | Access: Unauthenticated full CRUD")
    doc.add_paragraph("Discovered via $select[]=__INVALID__ — 26 fields:")

    user_fields = [
        ["_id", "ObjectId", "Low"],
        ["email", "String", "HIGH — PII"],
        ["phoneNumber", "String", "HIGH — PII"],
        ["password", "String", "CRITICAL — bcrypt hash"],
        ["googleId", "String", "HIGH — OAuth ID"],
        ["accountVerified", "Boolean", "Medium"],
        ["tempPhoneNumber", "String", "HIGH"],
        ["phoneOtp", "String", "CRITICAL — live OTP exposed in GET!"],
        ["phoneOtpValidTill", "String", "CRITICAL — OTP expiry"],
        ["passwordResetToken", "String", "CRITICAL — exposed in GET!"],
        ["passwordResetTokenExpiry", "String", "CRITICAL — token expiry"],
        ["dateOfBirth", "String", "HIGH — PII"],
        ["gender", "String", "Medium"],
        ["height", "Number", "Medium — health data"],
        ["weight", "Number", "Medium — health data"],
        ["createdAt", "String", "Low"],
        ["updatedAt", "String", "Low"],
        ["hasPremiumMembership", "Boolean", "Low"],
        ["rewardCoinsBalance", "Number", "Medium"],
        ["premiumMembership", "Object", "Medium"],
        ["profileToken", "String", "CRITICAL — session token"],
        ["profileTokenValidTill", "String", "HIGH"],
        ["identifierType", "String", "Low"],
        ["isWeb", "Boolean", "Low"],
        ["referralCode", "String", "Low"],
        ["referral", "Object", "Low"],
    ]
    add_table_from_data(doc, ["Field", "Type", "Sensitivity"], user_fields)

    add_callout(doc,
        "phoneOtp, passwordResetToken, and profileToken are returned in unauthenticated "
        "GET requests. This enables account takeover for ANY user.", RED
    )

    # Orders Schema
    doc.add_heading("Orders Collection", level=1)
    doc.add_paragraph("Total: ~63,000 records | Access: Any JWT (no RBAC)")
    doc.add_paragraph("108 total fields including nested objects. Populated references leak full user, address, and store data.")

    doc.add_heading("Top-Level Order Fields (33)", level=2)
    order_fields = [
        ["_id / orderId", "ObjectId / String", "Order identifiers"],
        ["status", "String", "paid, delivered, cancelled"],
        ["orderTotal / subTotal / paymentAmount", "Number", "Financial amounts"],
        ["currency", "String", "INR"],
        ["deliveryCharge / taxAmount / platformFee", "Number", "Fee breakdown"],
        ["paymentMode / paymentOrderId", "String", "payu, TXN[REDACTED]"],
        ["couponCode / offerId / discountedAmount", "Mixed", "Discount tracking"],
        ["hasPrescription / consultDoctorForPrescription", "Boolean", "Medical flags"],
        ["deliveryMode / deviceType / lastTimelineStatus", "String", "Order metadata"],
        ["isRewardCoinsApplied / rewardCoinsUsed", "Mixed", "Loyalty"],
        ["patientId", "ObjectId", "Patient reference (PHI link)"],
        ["createdAt / updatedAt", "String", "Timestamps"],
    ]
    add_table_from_data(doc, ["Field(s)", "Type", "Notes"], order_fields)

    doc.add_heading("Nested: items[] Array", level=2)
    add_table_from_data(doc, ["Field", "Type"],
        [["productId", "ObjectId"], ["quantity", "Number"], ["amount / total", "Number"], ["_id", "ObjectId"]])

    doc.add_heading("Nested: address Object (17 fields) — FULL USER ADDRESS", level=2)
    addr_fields = [
        ["userName", "String", "Full name"],
        ["phoneNumber / alternatePhoneNumber", "String", "Phone numbers"],
        ["addressLine1 / addressLine2 / fullAddress", "String", "Physical address"],
        ["postalCode / city / state / country", "String", "Location"],
        ["coordinates.latitude / .longitude", "Number", "GPS coordinates"],
        ["isDefault / type / userId", "Mixed", "Address metadata"],
    ]
    add_table_from_data(doc, ["Field(s)", "Type", "Notes"], addr_fields)

    doc.add_heading("Nested: store Object (21+ fields) — FULL STORE DATA", level=2)
    store_fields = [
        ["storeCode / storeName / storeId", "String", "Store identifiers"],
        ["gstNumber / fssaiNumber / licenceNumber", "String", "Regulatory — GST, FSSAI, drug licence"],
        ["email / phoneNumber", "String", "Contact info"],
        ["address / city / state / pincode / country", "String", "Physical location"],
        ["serviceableZip", "Array", "1,282 ZIP codes per store"],
        ["coordinates.latitude / .longitude", "Number", "GPS"],
        ["logistics.shiprocketQuick / .delhivery", "Object", "Logistics provider config"],
        ["active / acceptedInvitation", "Boolean", "Status flags"],
    ]
    add_table_from_data(doc, ["Field(s)", "Type", "Notes"], store_fields)

    doc.add_heading("Nested: userId Object (17 fields) — FULL USER POPULATED", level=2)
    doc.add_paragraph("Every order response includes the complete user document: email, phone, DOB, gender, OTP fields, profile tokens.")

    add_callout(doc,
        "A single GET /super-admin-users/orders?$limit=50 request returns 50 complete user records, "
        "50 full addresses with GPS, and 50 store configurations with GST/FSSAI numbers.", RED
    )

    # Attachments
    doc.add_heading("Attachments Collection", level=1)
    doc.add_paragraph("Total: 98,464 records | Access: Unauthenticated | Storage: AWS S3")
    attach_fields = [
        ["_id", "ObjectId", "Document ID"],
        ["storageService", "String", "S3"],
        ["objectDetails.size", "Number", "File size in bytes"],
        ["objectDetails.fileName", "String", "S3 key"],
        ["objectDetails.originalFileName", "String", "Original upload name"],
        ["objectDetails.mimeType", "String", "image/png, application/pdf, etc."],
        ["objectUrl", "String", "Full S3 URL — publicly accessible"],
    ]
    add_table_from_data(doc, ["Field", "Type", "Notes"], attach_fields)
    doc.add_paragraph("Includes prescription images, medical documents, and store assets — all publicly enumerable and downloadable.")

    # Zip Codes
    doc.add_heading("Zip Codes Collection", level=1)
    doc.add_paragraph("Total: 19,794 records | Access: Unauthenticated")
    zip_fields = [
        ["_id", "ObjectId"], ["area", "String"], ["district", "String"],
        ["state", "String"], ["zipCode", "String"],
        ["location.type", "String (Point)"], ["location.coordinates", "Array [lng, lat]"],
        ["isDeliverable", "Boolean"],
    ]
    add_table_from_data(doc, ["Field", "Type"], zip_fields)

    # Other collections
    doc.add_heading("Other Discovered Collections", level=1)

    doc.add_heading("app-data (7 fields)", level=2)
    doc.add_paragraph("Access: Unauthenticated | POST requires: name, type")
    add_table_from_data(doc, ["Field", "Type"],
        [["_id", "ObjectId"], ["name", "String"], ["visibility", "String"],
         ["type", "String"], ["description", "String"], ["imageUrl", "String"], ["value", "Mixed"]])

    doc.add_heading("application-tax (4 fields)", level=2)
    doc.add_paragraph("Access: Unauthenticated")
    add_table_from_data(doc, ["Field", "Type"],
        [["_id", "ObjectId"], ["defaultTax", "Number"], ["createdBy", "ObjectId"], ["updatedBy", "ObjectId"]])

    doc.add_heading("consultancy-appointment-slots (7 fields)", level=2)
    doc.add_paragraph("Access: Unauthenticated")
    add_table_from_data(doc, ["Field", "Type"],
        [["_id", "ObjectId"], ["date", "String"], ["startTime", "String"], ["endTime", "String"],
         ["maxAppointments", "Number"], ["appointments", "Array"], ["availableCount", "Number"]])

    doc.add_heading("policies (6 records)", level=2)
    doc.add_paragraph("Access: Unauthenticated | Returns raw array (no pagination)")
    add_table_from_data(doc, ["Field", "Type"],
        [["_id", "ObjectId"], ["type", "String"], ["name", "String"],
         ["value.url / .version / .lastUpdateAt", "Mixed"], ["content", "String (HTML)"]])

    doc.add_heading("patients (8 fields)", level=2)
    doc.add_paragraph("Access: JWT required | Returns raw array (no pagination)")
    add_table_from_data(doc, ["Field", "Type"],
        [["_id", "ObjectId"], ["name", "String"], ["phoneNumber", "String"],
         ["email", "String"], ["dateOfBirth", "String"], ["gender", "String"],
         ["relation", "String"], ["userId", "ObjectId"]])

    doc.add_heading("contact-us", level=2)
    doc.add_paragraph("Access: Unauthenticated read + write!")
    doc.add_paragraph("POST required fields: name -> email -> phoneNumber -> message")

    doc.add_heading("admin-zip-codes (2 fields)", level=2)
    add_table_from_data(doc, ["Field", "Type"], [["_id", "ObjectId"], ["text", "String"]])

    doc.add_heading("products (search endpoint)", level=2)
    doc.add_paragraph("Query params: id, category, sponsored, filter, page, limit, userId, cartSimilarProducts")

    # POST validation disclosure
    doc.add_heading("POST Validation Field Disclosure", level=1)
    doc.add_paragraph(
        "By sending empty {} POST requests and incrementally adding fields, we mapped required fields:"
    )
    post_chain = [
        ["contact-us", "name -> email -> phoneNumber -> message", "YES"],
        ["app-data", "name -> type", "YES"],
        ["payment", "No validation — creates records!", "YES (WRITE!)"],
        ["carts", "Requires JWT", "No"],
        ["consultations", "Requires JWT", "No"],
        ["memberships", "Requires JWT", "No"],
        ["patients", "Requires JWT", "No"],
    ]
    add_table_from_data(doc, ["Service", "Required Fields", "Unauthenticated?"], post_chain)

    add_callout(doc,
        "POST /payment accepts requests WITHOUT authentication and creates records. "
        "This is an unauthenticated write to a payment collection.", RED
    )

    # Mongoose Internals
    doc.add_heading("Mongoose Internals Leak", level=1)
    doc.add_paragraph("Endpoint: PATCH /users/{id} with empty body {}")
    doc.add_paragraph("Returns Mongoose document internals:")
    internals = [
        ["$__", "Internal Mongoose state object"],
        ["$isNew", "Document creation flag"],
        ["_doc", "Raw MongoDB document with ALL fields"],
    ]
    add_table_from_data(doc, ["Field", "Description"], internals)
    doc.add_paragraph(
        "The _doc field contains: password (bcrypt hash), phoneOtp (active OTP), "
        "passwordResetToken, profileToken, and all PII fields. "
        "This leak is unique to /users — other endpoints return normal JSON errors."
    )

    # Data Volume
    doc.add_heading("Data Volume Summary", level=1)
    volumes = [
        ["Users", "~188,000", "YES (full CRUD)"],
        ["Attachments", "98,464", "YES (read + S3 URLs)"],
        ["Orders", "~63,000", "JWT required (obtainable via OTP leak)"],
        ["Zip Codes", "19,794", "YES"],
        ["Medicine Requests", "~6,600", "JWT required"],
        ["Policies", "6", "YES"],
    ]
    add_table_from_data(doc, ["Collection", "Records", "Unauthenticated Access"], volumes)

    p = doc.add_paragraph()
    run = p.add_run("Estimated total PII records at risk: 250,000+")
    run.font.bold = True
    run.font.color.rgb = RED

    # Sensitive fields matrix
    doc.add_heading("Sensitive Field Exposure Matrix", level=1)
    sensitive = [
        ["Email addresses", "/users", "188,000"],
        ["Phone numbers", "/users, orders nested", "188,000+"],
        ["Physical addresses + GPS", "Orders -> address", "63,000"],
        ["Date of birth", "/users, orders -> userId", "188,000"],
        ["OTP codes", "/users -> phoneOtp", "Active OTPs"],
        ["Password reset tokens", "/users -> passwordResetToken", "Active tokens"],
        ["Password hashes", "/users via PATCH _doc", "188,000"],
        ["Medical prescriptions", "/attachments -> S3 URLs", "98,464 files"],
        ["Store GST/FSSAI/licence", "Orders -> store nested", "All stores"],
        ["Payment transaction IDs", "Orders -> paymentOrderId", "63,000"],
    ]
    add_table_from_data(doc, ["Data Type", "Source Endpoint", "Records"], sensitive)

    # Swagger Gap
    doc.add_heading("Swagger Gap Analysis", level=1)
    gaps = [
        ["Schema definitions", "All models", "0 definitions"],
        ["Security on endpoints", "All 567", "4 only (0.7%)"],
        ["Auth documentation", "Complete", "Minimal"],
        ["Rate limiting", "Documented", "None"],
        ["Input validation", "Documented", "None"],
    ]
    add_table_from_data(doc, ["Aspect", "Expected", "Actual"], gaps)

    # Store admin bugs
    doc.add_heading("Store Admin Endpoint Bugs", level=1)
    bugs = [
        ["POST /store-admin-users/forgot-password", "Cannot read properties of undefined (reading 'toLowerCase')"],
        ["POST /store-admin-users/reset-password", "Cannot read properties of undefined (reading 'toLowerCase')"],
    ]
    add_table_from_data(doc, ["Endpoint", "Error"], bugs)

    # Remediation
    doc.add_heading("Remediation Priority", level=1)

    doc.add_heading("IMMEDIATE (24-48 hours)", level=2)
    for item in [
        "Remove phoneOtp, passwordResetToken, profileToken, password from ALL API responses",
        "Add authentication to /users GET endpoint",
        "Add authentication to /attachments GET endpoint",
        "Disable unauthenticated POST to /payment",
        "Remove Mongoose internals from PATCH responses",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("HIGH (1 week)", level=2)
    for item in [
        "Implement per-user data scoping",
        "Add rate limiting to all endpoints",
        "Implement $limit ceiling (max 50, server-enforced)",
        "Remove nested population of sensitive data in order responses",
        "Fix store-admin-users error handling",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("MEDIUM (1 month)", level=2)
    for item in [
        "Add security definitions to all Swagger endpoints",
        "Implement proper RBAC across all 93 services",
        "Audit all 567 endpoint methods for authorization",
        "Implement field-level access control",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Signatures
    doc.add_heading("Signatures", level=1)
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Assessment conducted by:", ""),
        ("Date:", "2026-02-18"),
        ("Authorization reference:", ""),
        ("Client representative:", ""),
    ]):
        sig_table.rows[i].cells[0].text = label
        sig_table.rows[i].cells[1].text = val
        sig_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    path = os.path.join(OUT_DIR, "MedBasket_Database_Schema_Map.docx")
    doc.save(path)
    print(f"  Schema Map: {path}")
    return path


# ============================================================================
# DOCUMENT 5: Attachment Escalation Report
# ============================================================================
def build_attachment_escalation():
    doc = Document()
    style_doc(doc)

    add_cover_page(
        doc,
        "MedBasket API\nAttachment Storage Compromise",
        "Escalation: 98,476 Files Including Medical Prescriptions, Aadhaar Cards, and Government IDs",
    )

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph(
        "The /attachments API endpoint exposes 98,476 files without authentication, including "
        "1,897 medical prescriptions, 9 Aadhaar (national ID) cards, 4 PAN cards, 1 passport, "
        "79 hospital/medical reports, 59 customer invoices, 14 employee resumes, "
        "3,046 store inventory spreadsheets, and an estimated ~45,000 user-uploaded prescription photos."
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "4 Aadhaar card PDFs are DIRECTLY DOWNLOADABLE from the publicly accessible S3 development bucket. "
        "Aadhaar exposure is a criminal offense under Aadhaar Act 2016 Section 29(4)."
    )
    run.font.bold = True
    run.font.color.rgb = RED

    # Storage Infrastructure
    doc.add_heading("Storage Infrastructure", level=1)
    add_table_from_data(doc,
        ["Backend", "Domain/Bucket", "Est. Files", "Direct Download?"],
        [
            ["S3 Dev Bucket", "techpepo-development-s3.s3.ap-south-1.amazonaws.com", "~16,400", "YES — No auth"],
            ["Media CDN", "media.medbasket.com (CloudFront)", "~82,000", "No (HTTP 403)"],
        ]
    )

    doc.add_heading("S3 Dev Bucket Misconfiguration", level=2)
    s3_checks = [
        ["Directory listing", "AccessDenied (good)"],
        ["Object read", "PUBLIC — HTTP 200 for any file"],
        ["Object write (PUT)", "403 Forbidden (good)"],
        ["ACL read", "PUBLIC — AllUsers has READ_ACP"],
        ["AWS Owner ID", "[CANONICAL_ID_REDACTED]"],
    ]
    add_table_from_data(doc, ["Check", "Result"], s3_checks)

    add_callout(doc,
        "The S3 bucket ACL grants READ_ACP to AllUsers (anonymous). "
        "This is a development bucket (techpepo-development-s3) being used in PRODUCTION.", RED
    )

    # File Type Breakdown
    doc.add_heading("Complete File Type Breakdown", level=1)
    file_types = [
        ["JPEG images", "66,122", "Mostly CDN"],
        ["PNG images", "16,976", "Mixed"],
        ["JPG images", "5,749", "CDN"],
        ["PDF documents", "5,700", "Mixed"],
        ["XLSX spreadsheets", "3,046", "CDN"],
        ["HEIC/HEIF (iPhone)", "807", "CDN"],
        ["TOTAL", "98,476", "~274 GB estimated"],
    ]
    add_table_from_data(doc, ["Type", "Count", "Storage"], file_types)

    # Sensitive PDF Analysis
    doc.add_heading("Sensitive PDF Analysis (5,700 PDFs — Full Scan)", level=1)
    pdf_cats = [
        ["Medical Prescriptions", "1,897", "0", "1,897"],
        ["Scanned Documents", "481", "0", "481"],
        ["Medical/Hospital Reports", "79", "0", "79"],
        ["Medical Documents", "66", "0", "66"],
        ["Customer Invoices", "59", "0", "59"],
        ["Employee Resumes/CVs", "14", "1", "13"],
        ["Aadhaar Cards (National ID)", "9", "4", "5"],
        ["PAN Cards (Tax ID)", "4", "0", "4"],
        ["Other (passport, discharge, ECG/MRI)", "3", "0", "3"],
        ["Lab/Blood Reports", "2", "0", "2"],
        ["FSSAI License", "1", "0", "1"],
        ["TOTAL SENSITIVE", "2,615", "5", "2,610"],
    ]
    add_table_from_data(doc,
        ["Category", "Count", "S3 Dev (Downloadable)", "CDN (Metadata Leaked)"],
        pdf_cats
    )

    # Directly downloadable files
    doc.add_heading("Directly Downloadable Sensitive Files", level=1)
    p = doc.add_paragraph("These files are accessible by ANYONE by visiting the URL:")
    download_files = [
        ["aadhar.pdf", "970,774 bytes", "Aadhaar Card (National ID)"],
        ["aadhar.pdf", "970,774 bytes", "Aadhaar Card (National ID)"],
        ["aadhar.pdf", "970,774 bytes", "Aadhaar Card (National ID)"],
        ["aadhar.pdf", "970,774 bytes", "Aadhaar Card (National ID)"],
        ["[REDACTED_NAME]_Resume_9 (1).pdf", "132,503 bytes", "Employee Resume"],
    ]
    add_table_from_data(doc, ["Filename", "Size", "Document Type"], download_files)

    add_callout(doc,
        "Aadhaar card exposure is a CRIMINAL OFFENSE under Aadhaar Act 2016 Section 29(4). "
        "Penalty: imprisonment up to 3 years + fine up to Rs 10,000.", RED
    )

    # Prescription photos
    doc.add_heading("Prescription Photo Estimate (~72,000 images)", level=1)
    photo_patterns = [
        ["Numeric filenames (Android gallery)", "~30,700", "Prescription photos from app"],
        ["Descriptive names", "~19,500", "Mixed — includes prescriptions"],
        ["Phone camera (IMG_*, rn_image_picker_*)", "~10,600", "Mobile app prescription uploads"],
        ["WhatsApp images", "~3,600", "WhatsApp-shared prescriptions"],
    ]
    add_table_from_data(doc, ["Pattern", "Est. Count", "Likely Content"], photo_patterns)
    doc.add_paragraph(
        "Conservative estimate: ~45,000 of the 72,000 images are prescription photos. "
        "In a healthcare/pharmacy app, the majority of user-uploaded photos ARE prescriptions."
    )

    # Store Inventory
    doc.add_heading("Store Inventory Spreadsheets (3,046 XLSX)", level=1)
    xlsx_samples = [
        ["store-inventory (9).xlsx", "Store product inventory"],
        ["STOCK REPORT 6-1-25.xlsx", "Stock levels with dates"],
        ["BOMMANAHALLI STORE 08-01-25 E-com.xlsx", "Store e-commerce data"],
        ["store-inventory new.xlsx", "Updated inventory"],
    ]
    add_table_from_data(doc, ["Filename", "Content"], xlsx_samples)
    doc.add_paragraph("Contains business-confidential data: product lists, stock levels, pricing, store performance.")

    # Legal Impact
    doc.add_heading("Legal & Regulatory Impact", level=1)

    doc.add_heading("DPDP Act 2023", level=2)
    dpdp = [
        ["Section 8(4)", "Failure to implement reasonable security safeguards"],
        ["Section 12", "Mandatory breach notification to Data Protection Board"],
        ["Section 15", "Penalty up to Rs 250 Crore (~$30M USD) per instance"],
    ]
    add_table_from_data(doc, ["Section", "Violation"], dpdp)

    doc.add_heading("Aadhaar Act 2016", level=2)
    aadhaar = [
        ["Section 29(4)", "No person shall publicly display Aadhaar number"],
        ["Section 37", "Imprisonment up to 3 years + fine up to Rs 10,000"],
    ]
    add_table_from_data(doc, ["Section", "Violation"], aadhaar)

    doc.add_heading("Other Applicable Laws", level=2)
    other_laws = [
        ["IT Act Section 43A", "Body corporate must implement reasonable security for sensitive data"],
        ["IT Act Section 72A", "Disclosure of personal info — imprisonment up to 3 years"],
        ["Drugs & Cosmetics Act", "Prescription confidentiality violation (1,897+ prescriptions)"],
    ]
    add_table_from_data(doc, ["Law/Section", "Violation"], other_laws)

    # Impact Summary
    doc.add_heading("Impact Summary", level=1)
    impact = [
        ["Total files exposed", "98,476"],
        ["Files directly downloadable", "~16,400 (S3 dev bucket)"],
        ["Medical prescriptions (PDF)", "1,897"],
        ["Estimated prescription photos", "~45,000"],
        ["Government IDs (Aadhaar + PAN + Passport)", "14"],
        ["Hospital/medical reports", "145"],
        ["Customer invoices", "59"],
        ["Store inventory spreadsheets", "3,046"],
        ["Employee resumes", "14"],
        ["Estimated data volume", "~274 GB"],
        ["AWS account ID exposed", "Yes (via S3 ACL)"],
    ]
    add_table_from_data(doc, ["Metric", "Value"], impact)

    # Remediation
    doc.add_heading("Remediation", level=1)

    doc.add_heading("IMMEDIATE (within hours)", level=2)
    for item in [
        "Block unauthenticated access to /attachments — add authenticate('jwt') hook",
        "Remove public read from S3 dev bucket — set bucket policy to Block Public Access",
        "Remove READ_ACP grant for AllUsers on S3 dev bucket",
        "Delete or restrict Aadhaar/PAN/passport PDFs from S3 dev bucket",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 24 hours", level=2)
    for item in [
        "Implement per-user data scoping — users only see their own attachments",
        "Add rate limiting on attachments endpoint",
        "Migrate ALL files from dev bucket to production CDN with access controls",
        "Audit all stored government ID documents for compliance",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 week", level=2)
    for item in [
        "Implement signed URLs with short expiry for all file access",
        "Separate medical/prescription files with stricter access controls",
        "Data retention policy — delete files no longer needed",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Within 1 month", level=2)
    for item in [
        "Full data classification audit of all 98K+ files",
        "Implement encryption at rest for all sensitive files",
        "DPDP Act breach notification to Data Protection Board",
        "Patient notification about prescription exposure",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Signatures
    doc.add_heading("Signatures", level=1)
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Assessment conducted by:", ""),
        ("Date:", "2026-02-18"),
        ("Authorization reference:", ""),
        ("Client representative:", ""),
    ]):
        sig_table.rows[i].cells[0].text = label
        sig_table.rows[i].cells[1].text = val
        sig_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    path = os.path.join(OUT_DIR, "MedBasket_Attachment_Escalation.docx")
    doc.save(path)
    print(f"  Attachments:{path}")
    return path


# ============================================================================
# Main
# ============================================================================
def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    print("Generating DOCX files...")
    print()
    build_report()
    build_walkthrough()
    build_evidence()
    build_schema_map()
    build_attachment_escalation()
    print()
    print(f"All files saved to: {OUT_DIR}/")
    print("Ready to share with board members.")


if __name__ == "__main__":
    main()
